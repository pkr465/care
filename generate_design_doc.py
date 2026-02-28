#!/usr/bin/env python3
"""
Generate comprehensive CARE Design Document for Verilog/SystemVerilog HDL analysis.
Uses python-docx to create a professional Word document (.docx).
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

def set_cell_background(cell, fill):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_title_page(doc):
    """Add title page to document."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CARE')
    run.font.size = Pt(56)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Codebase Analysis & Repair Engine\nfor HDL (Verilog/SystemVerilog)')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(79, 129, 189)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Version 1.0')
    run.font.size = Pt(14)
    run.font.bold = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run('February 2026')
    run.font.size = Pt(14)

    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents."""
    toc_heading = doc.add_heading('Table of Contents', level=1)

    sections = [
        '1. Executive Summary',
        '2. Requirements Specification',
        '   2.1 Functional Requirements',
        '   2.2 Non-Functional Requirements',
        '3. Architecture Design',
        '   3.1 System Overview',
        '   3.2 Module Architecture',
        '   3.3 Database Architecture',
        '   3.4 LLM Integration',
        '   3.5 Scoring & Grading',
        '4. Configuration Reference',
        '5. Project File Structure',
        '6. Usage Guide',
        '7. API & Interface Specification',
        '8. Output Formats',
        '9. Extension Guide',
        '10. Dependencies & Technology Stack',
    ]

    for section in sections:
        if section.startswith('   '):
            p = doc.add_paragraph(section[3:], style='List Bullet 2')
        elif section.startswith('2.') or section.startswith('3.'):
            p = doc.add_paragraph(section, style='List Number')
        else:
            p = doc.add_paragraph(section, style='List Number')

    doc.add_page_break()

def add_executive_summary(doc):
    """Add Executive Summary section."""
    doc.add_heading('1. Executive Summary', level=1)

    intro = """CARE (Codebase Analysis & Repair Engine) is a comprehensive framework for automated analysis, design review, and repair of Verilog/SystemVerilog HDL codebases. Originally derived from the CURE architecture, CARE adapts and extends these capabilities specifically for hardware design languages, enabling teams to assess and improve digital design quality at scale."""

    doc.add_paragraph(intro)

    doc.add_heading('1.1 Overview', level=2)
    overview_points = [
        'Multi-Stage Analysis Pipeline: Combines static analysis, deep integration with HDL tools (Verilator, Verible), and LLM-powered design review',
        'Nine Specialized Analyzers: Quality, Complexity, Synthesis Safety, Documentation, Maintainability, Verification Coverage, Clock Domain Crossing (CDC), Uninitialized Signal Detection, and Signal Integrity',
        'Five Deep Analysis Adapters: HDL Complexity (Verilator), Lint (Veriblle/Verilator), Unused Modules, Hierarchy Analysis, and Module-Level Metrics',
        'Multi-Provider LLM Integration: Support for Anthropic Claude, QGenie, Google Vertex AI, and Azure OpenAI',
        'Design Health Scoring: Weighted quantitative assessment across eight dimensions',
        'Synthesis Safety Detection: Identifies CDC violations, combinational loops, metastability risks, and latch inference issues',
        'Automated HDL Repair: LLM-powered suggestions and automated refactoring for common design issues',
        'Human-in-the-Loop (HITL): Persistent feedback storage and RAG-based constraint injection for custom design rules',
        'Vector Database Integration: PostgreSQL with pgvector for semantic search and knowledge retention',
        'Comprehensive Reporting: HTML health reports, Excel design reviews, JSON data export, and natural language summaries'
    ]

    for point in overview_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading('1.2 Key Capabilities', level=2)

    table = doc.add_table(rows=10, cols=2)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Capability'
    header_cells[1].text = 'Description'

    set_cell_background(header_cells[0], '4F81BD')
    set_cell_background(header_cells[1], '4F81BD')
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    capabilities = [
        ('Static Analysis', 'Pattern-based detection of coding style violations, unsafe constructs, and design anti-patterns'),
        ('Tool Integration', 'Deep integration with Verilator and Verible for syntax validation and linting'),
        ('LLM Review', 'Claude-powered design review providing architectural feedback and refactoring suggestions'),
        ('Hierarchy Extraction', 'Module instantiation graph construction with cross-reference analysis'),
        ('Testability Assessment', 'Verification coverage estimation and UVM/SVA pattern detection'),
        ('Design Metrics', 'Cyclomatic complexity, HDL-specific metrics, and design health scoring'),
        ('Safety Analysis', 'Synthesis safety checks including CDC, combinational loops, and timing constraints'),
        ('Report Generation', 'Multi-format output: HTML reports, Excel sheets, JSON data, and natural language summaries'),
        ('Vector Database', 'Semantic search across codebase with RAG-powered knowledge augmentation'),
    ]

    for i, (capability, description) in enumerate(capabilities, 1):
        row = table.rows[i]
        row.cells[0].text = capability
        row.cells[1].text = description

    doc.add_page_break()

def add_requirements_specification(doc):
    """Add Requirements Specification section."""
    doc.add_heading('2. Requirements Specification', level=1)

    doc.add_heading('2.1 Functional Requirements', level=2)

    doc.add_heading('2.1.1 Static Analysis Engine', level=3)
    req_text = """The system shall implement nine specialized static analyzers for Verilog/SystemVerilog HDL codebases:

• QualityAnalyzer: Detects HDL coding quality violations including blocking assignment misuse, non-blocking assignment in combinational logic, improper latch inference, missing sensitivity lists, and race conditions.

• ComplexityAnalyzer: Measures always block complexity, module complexity, nesting depth, and cyclomatic complexity metrics. Calculates HDL-specific complexity indexes.

• SynthesisSafetyAnalyzer: Identifies synthesis-unsafe constructs including combinational loops, asynchronous reset violations, CDC (Clock Domain Crossing) hazards, metastability risks, and timing constraint violations.

• DocumentationAnalyzer: Measures HDL documentation coverage including module documentation, port documentation, always block comments, and parameter/localparam documentation.

• MaintainabilityAnalyzer: Calculates maintainability index based on cyclomatic complexity, lines of code, Halstead metrics, and comment density adjusted for HDL-specific patterns.

• VerificationCoverageAnalyzer: Detects testbench presence, UVM framework usage, SVA (SystemVerilog Assertion) patterns, coverage model usage, and functional coverage metrics.

• CDCAnalyzer: Specialized clock domain crossing analysis detecting CDC violations, gray code encoder/decoder pairs, synchronizer usage, and metastability protection.

• UninitializedSignalAnalyzer: Identifies undriven signals, floating ports, uninitialized variables, and incomplete case statements with default branches.

• SignalIntegrityAnalyzer: Detects bus contention, port width mismatches, signed/unsigned comparison issues, and illegal type conversions."""

    doc.add_paragraph(req_text)

    doc.add_heading('2.1.2 Deep Analysis Adapters', level=3)
    adapter_text = """The system shall integrate external HDL analysis tools through five specialized adapters:

• HDLComplexityAdapter: Uses Verilator's hierarchical analysis to extract module complexity metrics, port statistics, parameter usage, and instantiation depth.

• LintAdapter: Integrates Verible and Verilator linting providing systematic error/warning categorization, severity classification, and rule-based filtering.

• UnusedModuleAdapter: Analyzes module instantiation hierarchy to identify unused modules, unreachable code paths, and orphaned design blocks.

• HierarchyAnalyzerAdapter: Constructs module instantiation graphs, tracks port connectivity, detects hierarchical structure issues, and analyzes design topology.

• ModuleMetricsAdapter: Collects per-module metrics including input/output counts, parameter counts, instantiation count, and design footprint."""

    doc.add_paragraph(adapter_text)

    doc.add_heading('2.1.3 LLM HDL Review', level=3)
    llm_text = """The system shall integrate large language models for semantic code analysis and design review:

• Smart HDL Chunking: Intelligently partition HDL modules into LLM-consumable chunks respecting module boundaries, preserving context, and maintaining semantic coherence.

• Context Injection: Automatically inject parameter definitions, macro expansions, include file context, and cross-module references to reduce false positives.

• Design Review: Generate architectural feedback, refactoring suggestions, style improvements, and design pattern recommendations based on HDL semantics.

• Multi-Provider Support: Support Anthropic Claude (primary), QGenie, Google Vertex AI, and Azure OpenAI with provider-agnostic abstraction layer.

• Reasoning Integration: Leverage extended reasoning capabilities for complex design rule verification and architectural constraint analysis."""

    doc.add_paragraph(llm_text)

    doc.add_heading('2.1.4 Human-in-the-Loop (HITL)', level=3)
    hitl_text = """The system shall provide human-in-the-loop feedback mechanisms for continuous improvement:

• Persistent Feedback Storage: Store analyst feedback on code issues in PostgreSQL-backed database.

• RAG-Based Constraint Injection: Use Retrieval-Augmented Generation to inject similar past cases and custom design rules into LLM prompts.

• Feedback-Driven Refinement: Iteratively improve analysis accuracy based on analyst corrections and feedback.

• Design Rule Management: Support custom design rule definition and application across analyses."""

    doc.add_paragraph(hitl_text)

    doc.add_heading('2.1.5 Automated HDL Repair', level=3)
    repair_text = """The system shall provide automated repair suggestions and implementations:

• Issue Classification: Categorize detected issues by severity, fixability, and impact.

• Repair Suggestions: Generate LLM-powered refactoring suggestions with before/after code examples.

• Automated Fixes: Implement automated fixes for common issues including style normalization, latch inference prevention, and sensitivity list correction.

• Safe Refactoring: Ensure refactored code maintains functional equivalence with original through verification checks."""

    doc.add_paragraph(repair_text)

    doc.add_heading('2.1.6 Vector Database & Semantic Search', level=3)
    vector_text = """The system shall implement semantic search and knowledge retention:

• Vector Embeddings: Convert HDL code chunks into semantic embeddings for similarity search.

• PostgreSQL + pgvector: Store embeddings in PostgreSQL with vector distance search capabilities.

• Semantic RAG: Use retrieval-augmented generation for context-aware analysis and issue resolution.

• Knowledge Retention: Maintain searchable history of analyses, issues, and resolutions."""

    doc.add_paragraph(vector_text)

    doc.add_heading('2.1.7 Dashboard & Reporting', level=3)
    report_text = """The system shall provide comprehensive reporting and visualization:

• HTML Health Reports: Interactive HTML reports with design health scores, issue breakdowns, and severity visualizations.

• Excel Analysis: Structured Excel workbooks with issue details, fix recommendations, and designer feedback.

• JSON Export: Machine-readable JSON export for automated tool integration.

• Natural Language Summaries: LLM-generated executive summaries of analysis results.

• Trend Analysis: Track design quality metrics over time and across iterations."""

    doc.add_paragraph(report_text)

    doc.add_heading('2.2 Non-Functional Requirements', level=2)

    nfr = [
        ('Performance', 'Analyze large codebases (>100k lines) within 5 minutes on standard hardware'),
        ('Scalability', 'Support incremental analysis, caching, and parallel processing'),
        ('Reliability', '99.5% uptime for analysis services with graceful degradation'),
        ('Maintainability', 'Modular architecture enabling independent analyzer updates'),
        ('Extensibility', 'Plugin architecture for custom analyzers and adapters'),
        ('Security', 'No hardcoded credentials, support for environment variables and secure vaults'),
        ('Portability', 'Cross-platform support (Linux, macOS, Windows with WSL)'),
        ('Configuration', 'Comprehensive YAML configuration with environment variable overrides'),
        ('Logging', 'Detailed logging with DEBUG, INFO, WARNING, ERROR levels'),
        ('Documentation', 'Complete API documentation, usage guides, and examples'),
    ]

    table = doc.add_table(rows=len(nfr) + 1, cols=2)
    table.style = 'Light Grid Accent 1'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Requirement'
    header_cells[1].text = 'Specification'

    set_cell_background(header_cells[0], '4F81BD')
    set_cell_background(header_cells[1], '4F81BD')
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    for i, (req, spec) in enumerate(nfr, 1):
        row = table.rows[i]
        row.cells[0].text = req
        row.cells[1].text = spec

    doc.add_page_break()

def add_architecture_design(doc):
    """Add Architecture Design section."""
    doc.add_heading('3. Architecture Design', level=1)

    doc.add_heading('3.1 System Overview', level=2)
    overview = """CARE follows a modular, agent-based architecture with clear separation of concerns. The system comprises five primary agents coordinated by a workflow orchestrator:

1. StaticAnalyzerAgent: Executes all nine static analyzers in parallel, computing foundational metrics.
2. DeepAnalysisAgent: Runs optional adapters (Verilator, Verible integration) for tool-specific insights.
3. LLMReviewAgent: Invokes language models for semantic analysis and design review.
4. MetricsAggregatorAgent: Consolidates metrics into unified health scores and weighted assessments.
5. ReportGeneratorAgent: Produces HTML, Excel, JSON, and natural language outputs.

All agents implement a common interface with consistent error handling and graceful degradation. The system operates in phases:

Phase 1 - Analysis: Parallel execution of static analyzers and optional deep adapters
Phase 2 - LLM Review: Chunked code submission to language models with context injection
Phase 3 - Aggregation: Consolidation into health metrics and design scores
Phase 4 - Reporting: Multi-format output generation"""

    doc.add_paragraph(overview)

    doc.add_heading('3.2 Module Architecture', level=2)

    arch_text = """The codebase is organized into the following modular structure:"""
    doc.add_paragraph(arch_text)

    modules = [
        ('agents/', 'Core agent implementations'),
        ('agents/core/', 'MetricsCalculator, analysis coordinators'),
        ('agents/analyzers/', 'Nine HDL-specific analyzer implementations'),
        ('agents/adapters/', 'Tool integration adapters (Verilator, Verible)'),
        ('agents/context/', 'Context builders for parameter/macro injection'),
        ('agents/parsers/', 'Input parsing and output generation'),
        ('utils/common/', 'LLM provider abstraction, logging utilities'),
        ('utils/database/', 'PostgreSQL and pgvector integration'),
        ('utils/hdl/', 'HDL-specific utilities and pattern matchers'),
    ]

    for path, desc in modules:
        doc.add_paragraph(f'{path}: {desc}', style='List Bullet')

    doc.add_heading('3.2.1 Analyzer Implementations', level=3)

    analyzer_table = doc.add_table(rows=10, cols=3)
    analyzer_table.style = 'Light Grid Accent 1'

    header_cells = analyzer_table.rows[0].cells
    header_cells[0].text = 'Analyzer'
    header_cells[1].text = 'File'
    header_cells[2].text = 'Purpose'

    for cell in header_cells:
        set_cell_background(cell, '4F81BD')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    analyzers = [
        ('QualityAnalyzer', 'quality_analyzer.py', 'HDL coding quality violations'),
        ('ComplexityAnalyzer', 'complexity_analyzer.py', 'Always block and module complexity'),
        ('SynthesisSafetyAnalyzer', 'security_analyzer.py', 'Synthesis safety hazards'),
        ('DocumentationAnalyzer', 'documentation_analyzer.py', 'Documentation coverage'),
        ('MaintainabilityAnalyzer', 'maintainability_analyzer.py', 'Maintainability index'),
        ('VerificationCoverageAnalyzer', 'test_coverage_analyzer.py', 'Testbench and coverage metrics'),
        ('CDCAnalyzer', 'potential_deadlock_analyzer.py', 'Clock domain crossing'),
        ('UninitializedSignalAnalyzer', 'null_pointer_analyzer.py', 'Undriven signal detection'),
        ('SignalIntegrityAnalyzer', 'memory_corruption_analyzer.py', 'Bus contention and type mismatches'),
    ]

    for i, (analyzer, file, purpose) in enumerate(analyzers, 1):
        row = analyzer_table.rows[i]
        row.cells[0].text = analyzer
        row.cells[1].text = file
        row.cells[2].text = purpose

    doc.add_heading('3.2.2 Adapter Implementations', level=3)

    adapter_table = doc.add_table(rows=6, cols=3)
    adapter_table.style = 'Light Grid Accent 1'

    header_cells = adapter_table.rows[0].cells
    header_cells[0].text = 'Adapter'
    header_cells[1].text = 'File'
    header_cells[2].text = 'Purpose'

    for cell in header_cells:
        set_cell_background(cell, '4F81BD')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    adapters = [
        ('HDLComplexityAdapter', 'hierarchy_complexity_adapter.py', 'Verilator-based complexity'),
        ('LintAdapter', 'design_rule_adapter.py', 'Verible/Verilator linting'),
        ('UnusedModuleAdapter', 'unused_signal_adapter.py', 'Unused module detection'),
        ('HierarchyAnalyzerAdapter', 'module_graph_adapter.py', 'Hierarchy graph construction'),
        ('ModuleMetricsAdapter', 'port_metrics_adapter.py', 'Module-level metrics'),
    ]

    for i, (adapter, file, purpose) in enumerate(adapters, 1):
        row = adapter_table.rows[i]
        row.cells[0].text = adapter
        row.cells[1].text = file
        row.cells[2].text = purpose

    doc.add_heading('3.3 Database Architecture', level=2)

    db_text = """CARE uses PostgreSQL with pgvector extension for both persistent storage and semantic search:

• Primary Storage: Analyzed design metrics, hierarchies, and issue databases in PostgreSQL tables
• Vector Embeddings: HDL code chunks embedded into semantic vectors via OpenAI/Veribk embedding API
• Semantic Search: pgvector extension enables L2 distance search for similar code patterns
• HITL Feedback: Persistent storage of analyst feedback and custom design rules
• Telemetry: Usage statistics and analysis metrics for framework improvement

Connection Configuration:
- Connection String: postgresql+psycopg2://user:password@host:port/database
- Pool Size: 5 persistent connections (configurable)
- Pool Recycling: 3600 seconds (auto-reset stale connections)
- SSL/TLS: Configurable for remote database servers"""

    doc.add_paragraph(db_text)

    doc.add_heading('3.4 LLM Integration', level=2)

    llm_text = """CARE provides abstraction over multiple LLM providers:

Provider Support:
• Anthropic Claude (primary): Via anthropic SDK with extended reasoning
• QGenie: Local/cloud deployment with LLaMA and Qwen models
• Google Vertex AI: Gemini models with enterprise support
• Azure OpenAI: Enterprise deployment option

Smart HDL Chunking:
The system respects HDL semantics when chunking code:
• Module-respecting boundaries: Never split within module definitions
• Context preservation: Include necessary `include files, parameters, typedefs
• Semantic coherence: Group related always blocks and instantiations
• Size limits: Target ~2000 tokens per chunk for LLM efficiency

Context Injection:
For each chunk, automatically include:
• Parameter and localparam definitions
• Macro definitions from `include files
• Port definitions and signal declarations
• Clock domain and reset strategy context
• Synthesis target and timing constraints

Prompt Engineering:
• Domain-specific system prompts for HDL design review
• Few-shot examples for pattern recognition
• Chain-of-thought reasoning for complex violations
• Structured output formats (JSON) for automated processing"""

    doc.add_paragraph(llm_text)

    doc.add_heading('3.5 Scoring & Grading', level=2)

    scoring_text = """CARE computes design health through weighted aggregation of eight dimensions:"""

    doc.add_paragraph(scoring_text)

    scoring_table = doc.add_table(rows=9, cols=3)
    scoring_table.style = 'Light Grid Accent 1'

    header_cells = scoring_table.rows[0].cells
    header_cells[0].text = 'Dimension'
    header_cells[1].text = 'Weight'
    header_cells[2].text = 'Description'

    for cell in header_cells:
        set_cell_background(cell, '4F81BD')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    scores = [
        ('Synthesis Risk', '25%', 'CDC, combinational loops, timing violations'),
        ('Lint Score', '20%', 'Verible/Verilator linting results'),
        ('Quality', '15%', 'Coding style, safe constructs, anti-patterns'),
        ('Maintainability', '15%', 'Complexity, documentation, readability'),
        ('Complexity', '10%', 'Cyclomatic and HDL-specific complexity'),
        ('Hierarchy', '5%', 'Module structure and reusability'),
        ('Documentation', '5%', 'Comment and documentation coverage'),
        ('Verification Coverage', '5%', 'Testbench presence and coverage metrics'),
    ]

    for i, (dim, weight, desc) in enumerate(scores, 1):
        row = scoring_table.rows[i]
        row.cells[0].text = dim
        row.cells[1].text = weight
        row.cells[2].text = desc

    doc.add_paragraph()
    scoring_formula = """Overall Design Health Score =
    (synthesis_risk × 0.25) + (lint_score × 0.20) + (quality × 0.15) +
    (maintainability × 0.15) + (complexity × 0.10) + (hierarchy × 0.05) +
    (documentation × 0.05) + (verification_coverage × 0.05)

Score Interpretation:
• 90-100: Excellent design quality, minimal issues
• 80-89: Good quality, minor improvements recommended
• 70-79: Acceptable quality, moderate improvements needed
• 60-69: Poor quality, significant improvements required
• <60: Critical issues, redesign recommended"""

    doc.add_paragraph(scoring_formula, style='List Bullet')

    doc.add_page_break()

def add_configuration_reference(doc):
    """Add Configuration Reference section."""
    doc.add_heading('4. Configuration Reference', level=1)

    config_text = """CARE is configured via global_config.yaml in the project root. All values support environment variable overrides using ${ENV_VAR} syntax.

Core Configuration Sections:"""

    doc.add_paragraph(config_text)

    doc.add_heading('4.1 Paths', level=2)
    doc.add_paragraph("""
source_dir: ./rtl                           # RTL source directory
code_base_path: ./rtl                       # Codebase path (same as source_dir)
out_dir: ./out                              # Output directory for reports
generated_md_dir: ./out/md                  # Markdown output
flat_json_path: ./out/parseddata            # Flattened JSON data
graph_path: ./out/diagrams                  # Diagram outputs
pdf_path: ./out/pdfs                        # PDF reports
prompt_file_path: ./prompts/prompt.md       # LLM prompt template
""")

    doc.add_heading('4.2 LLM Configuration', level=2)
    doc.add_paragraph("""
llm_provider: "anthropic"                   # anthropic | qgenie | vertexai | azure
model: "anthropic::claude-sonnet-4-20250514" # Primary analysis model
coding_model: "anthropic::claude-sonnet-4-20250514" # Code refactoring model
max_tokens: 16384                           # Maximum LLM output tokens
temperature: 0.1                            # Creativity (0.0 = deterministic)
timeout: 120                                # Request timeout in seconds
max_retries: 2                              # Retry attempts on failure
max_prompt_tokens: 100000                   # Token budget for prompt truncation
""")

    doc.add_heading('4.3 Database Configuration', level=2)
    doc.add_paragraph("""
connection: postgresql+psycopg2://...       # SQLAlchemy connection string
host: localhost                             # Database host
port: 5432                                  # Database port
database: codebase_analytics_db             # Database name
username: codebase_analytics_user           # Application user
password: postgres                         # Password (set in global_config.yaml)
pool_size: 5                                # Connection pool size
pool_recycle: 3600                          # Connection recycle timeout
ssl_mode: prefer                            # SSL/TLS mode
""")

    doc.add_heading('4.4 Scanning Configuration', level=2)
    doc.add_paragraph("""
exclude_dirs:                               # Directories to skip
  - sim_results
  - synthesis
  - .Xil
  - work
  - xsim.dir

exclude_globs:                              # File patterns to skip
  - "*.vcd"
  - "*.wlf"
  - "*/sim/*"
  - "*/generated/*"
""")

    doc.add_heading('4.5 HDL Tool Configuration', level=2)
    doc.add_paragraph("""
eda_tools:
  verilator_path: verilator                 # Verilator executable
  verible_syntax_path: verible-verilog-syntax
  verible_lint_path: verible-verilog-lint
  iverilog_path: iverilog                   # Icarus Verilog (optional)
  yosys_path: yosys                         # Yosys synthesis (optional)

hierarchy_builder:
  verible_executable: verible-verilog-syntax
  verilator_executable: verilator
  indexing_timeout_seconds: 300
  cache_metadata_filename: .cache_metadata.json
""")

    doc.add_heading('4.6 Context Injection Configuration', level=2)
    doc.add_paragraph("""
context:
  enable_include_context: true              # Enable parameter/macro injection
  hdl_include_paths: []                     # Additional -I style paths
  max_include_depth: 2                      # Follow `include chains depth
  max_context_chars: 6000                   # Max context per chunk
  exclude_includes: []                      # Exclude certain include files
""")

    doc.add_heading('4.7 HITL Configuration', level=2)
    doc.add_paragraph("""
hitl:
  enable: false                             # Enable human-in-the-loop
  rag_top_k: 5                              # Retrieve top 5 similar cases
  rag_similarity_threshold: 0.6             # Similarity threshold
  enable_prompt_augmentation: true          # Inject HITL feedback into prompts
  rag_context_max_tokens: 2000              # Max context from RAG
""")

    doc.add_page_break()

def add_project_structure(doc):
    """Add Project File Structure section."""
    doc.add_heading('5. Project File Structure', level=1)

    structure = """CARE/
├── main.py                             # Entry point
├── global_config.yaml                  # Configuration file
├── agents/
│   ├── core/
│   │   ├── metrics_calculator.py       # Metrics orchestrator
│   │   ├── analyzer_executor.py        # Parallel analyzer execution
│   │   └── aggregator.py               # Score aggregation
│   ├── analyzers/
│   │   ├── quality_analyzer.py         # HDL quality checks
│   │   ├── complexity_analyzer.py      # Complexity metrics
│   │   ├── security_analyzer.py        # Synthesis safety
│   │   ├── documentation_analyzer.py   # Documentation coverage
│   │   ├── maintainability_analyzer.py # Maintainability index
│   │   ├── test_coverage_analyzer.py   # Verification coverage
│   │   ├── potential_deadlock_analyzer.py # CDC analysis
│   │   ├── null_pointer_analyzer.py    # Uninitialized signals
│   │   └── memory_corruption_analyzer.py # Signal integrity
│   ├── adapters/
│   │   ├── base_adapter.py             # Base adapter interface
│   │   ├── ast_complexity_adapter.py   # Verilator AST complexity
│   │   ├── call_graph_adapter.py       # Module instantiation graph
│   │   ├── dead_code_adapter.py        # Dead signal detection
│   │   ├── function_metrics_adapter.py # Port/parameter metrics
│   │   ├── security_adapter.py         # Security pattern analysis
│   │   └── excel_report_adapter.py     # Excel generation
│   ├── context/
│   │   └── header_context_builder.py   # Include resolution, macro parsing
│   ├── reports/
│   │   └── complexity_report_pdf.py    # PDF complexity report
│   └── parsers/
│       ├── healthreport_generator.py   # HTML report generation
│       ├── json_flattener.py           # JSON flattening
│       └── ndjson_processor.py         # NDJSON conversion
├── utils/
│   ├── common/
│   │   ├── llm_tools.py                # LLM provider abstraction
│   │   ├── llm_tools_anthropic.py      # Anthropic integration
│   │   ├── llm_tools_qgenie.py         # QGenie integration
│   │   └── logging_utils.py            # Logging configuration
│   ├── database/
│   │   ├── postgres_manager.py         # PostgreSQL operations
│   │   ├── vector_db_pipeline.py       # pgvector integration
│   │   └── migrations.py               # Database migrations
│   ├── hdl/
│   │   ├── pattern_matcher.py          # Regex patterns
│   │   ├── verilog_parser.py           # Verilog parsing
│   │   └── systemverilog_extensions.py # SV-specific parsing
│   └── parsers/
│       └── global_config_parser.py     # YAML config parsing
├── prompts/
│   ├── prompt.md                       # HDL analysis prompt
│   └── chat_prompt.md                  # Chat mode prompt
├── data/
│   ├── docs/                           # Input documentation
│   └── examples/                       # Example HDL files
├── out/
│   ├── md/                             # Markdown outputs
│   ├── parseddata/                     # JSON analysis results
│   ├── diagrams/                       # Generated diagrams
│   └── pdfs/                           # PDF reports
├── tests/
│   ├── test_analyzers.py               # Analyzer unit tests
│   ├── test_adapters.py                # Adapter integration tests
│   └── test_integration.py             # End-to-end tests
├── requirements.txt                    # Python dependencies
├── setup.py                            # Installation script
└── README.md                           # Documentation"""

    doc.add_paragraph(structure, style='List Paragraph')

    doc.add_page_break()

def add_usage_guide(doc):
    """Add Usage Guide section."""
    doc.add_heading('6. Usage Guide', level=1)

    doc.add_heading('6.1 Installation', level=2)

    installation = """Clone the repository:
git clone <repository-url>
cd CARE

Install Python dependencies:
pip install -r requirements.txt

Install HDL tools (optional but recommended):
# Ubuntu/Debian
sudo apt-get install verilator verible

# macOS
brew install verilator verible

# Windows (requires WSL)
wsl apt-get install verilator verible"""

    doc.add_paragraph(installation, style='List Paragraph')

    doc.add_heading('6.2 Configuration', level=2)

    config = """1. Copy the default configuration:
cp global_config.yaml.example global_config.yaml

2. Edit global_config.yaml to set:
   - paths.source_dir: Path to your RTL code
   - paths.out_dir: Output directory for reports
   - llm.model: Your preferred LLM (default: Claude)
   - database.password: Set directly in global_config.yaml (if using vector DB)
   - email.smtp_password: Set directly in global_config.yaml (if using email)

3. API keys (.env file — only LLM keys go here):
LLM_API_KEY="sk-..."
QGENIE_API_KEY="..." (optional)"""

    doc.add_paragraph(config, style='List Paragraph')

    doc.add_heading('6.3 Basic Usage', level=2)

    basic = """Analyze a Verilog/SystemVerilog codebase:

python main.py --codebase-path ./rtl --output-dir ./out

With specific options:
python main.py \\
  --codebase-path ./rtl \\
  --output-dir ./out \\
  --enable-adapters \\
  --enable-llm \\
  --enable-vector-db

Generate HTML report:
python main.py --codebase-path ./rtl --generate-report

Analyze with all features enabled:
python main.py --codebase-path ./rtl --full-analysis"""

    doc.add_paragraph(basic, style='List Paragraph')

    doc.add_heading('6.4 Advanced Usage', level=2)

    advanced_sections = [
        ('HITL Integration', """Enable human-in-the-loop feedback:
python main.py --enable-hitl --feedback-file feedback.xlsx

This reads analyst feedback from feedback.xlsx and injects similar cases into analysis."""),

        ('Vector Database', """Enable semantic search:
python main.py --enable-vector-db

Requires PostgreSQL with pgvector extension configured in global_config.yaml"""),

        ('Incremental Analysis', """Analyze only changed files:
python main.py --incremental --last-run-timestamp 2026-02-27T00:00:00"""),

        ('Parallel Processing', """Use all available CPU cores:
python main.py --max-workers 8 --codebase-path ./rtl"""),

        ('Custom Analyzers', """Enable only specific analyzers:
python main.py --enable-analyzers quality,complexity,synthesis_safety"""),

        ('Tool Integration', """Use specific HDL tools:
python main.py --verilator-path /usr/bin/verilator \\
  --verible-path /usr/bin/verible-verilog-lint"""),
    ]

    for title, content in advanced_sections:
        doc.add_heading(f'6.4.{title}', level=3)
        doc.add_paragraph(content, style='List Paragraph')

    doc.add_page_break()

def add_api_specification(doc):
    """Add API & Interface Specification section."""
    doc.add_heading('7. API & Interface Specification', level=1)

    doc.add_heading('7.1 CLI Arguments', level=2)

    cli_table = doc.add_table(rows=16, cols=3)
    cli_table.style = 'Light Grid Accent 1'

    header_cells = cli_table.rows[0].cells
    header_cells[0].text = 'Argument'
    header_cells[1].text = 'Type'
    header_cells[2].text = 'Description'

    for cell in header_cells:
        set_cell_background(cell, '4F81BD')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    cli_args = [
        ('--codebase-path PATH', 'String', 'Path to RTL source code'),
        ('--output-dir PATH', 'String', 'Output directory for reports'),
        ('--enable-adapters', 'Flag', 'Enable deep analysis adapters'),
        ('--enable-llm', 'Flag', 'Enable LLM design review'),
        ('--enable-vector-db', 'Flag', 'Enable vector database integration'),
        ('--enable-hitl', 'Flag', 'Enable human-in-the-loop feedback'),
        ('--generate-report', 'Flag', 'Generate HTML health report'),
        ('--full-analysis', 'Flag', 'Run all analysis features'),
        ('--exclude-dirs DIRS', 'String', 'Comma-separated directories to exclude'),
        ('--exclude-globs PATTERNS', 'String', 'Comma-separated glob patterns to exclude'),
        ('--max-workers N', 'Integer', 'Number of parallel workers'),
        ('--debug', 'Flag', 'Enable debug logging'),
        ('--config PATH', 'String', 'Path to global_config.yaml'),
        ('--version', 'Flag', 'Print version and exit'),
        ('--help', 'Flag', 'Print help message'),
    ]

    for i, (arg, type_, desc) in enumerate(cli_args, 1):
        row = cli_table.rows[i]
        row.cells[0].text = arg
        row.cells[1].text = type_
        row.cells[2].text = desc

    doc.add_heading('7.2 Python API', level=2)

    api_text = """MetricsCalculator Class:

from agents.core.metrics_calculator import MetricsCalculator

calculator = MetricsCalculator(
    codebase_path='./rtl',
    output_dir='./out',
    enable_adapters=True,
    debug=False
)

metrics = calculator.compute_all_metrics()

Key Methods:
• compute_all_metrics(): Run all analyzers and return unified metrics dict
• compute_quality_metrics(): Run QualityAnalyzer only
• compute_complexity_metrics(): Run ComplexityAnalyzer only
• compute_synthesis_safety_metrics(): Run SynthesisSafetyAnalyzer only
• get_design_health_score(): Calculate overall design health (0-100)
• export_json(filepath): Export metrics to JSON
• export_excel(filepath): Export metrics to Excel"""

    doc.add_paragraph(api_text)

    doc.add_heading('7.3 Analyzer Interface', level=2)

    analyzer_interface = """All analyzers implement common interface:

class BaseAnalyzer:
    def analyze(self, file_path: str) -> Dict[str, Any]:
        '''Analyze single file and return metrics dict'''
        pass

    def batch_analyze(self, file_paths: List[str]) -> List[Dict]:
        '''Analyze multiple files in parallel'''
        pass

    def get_score(self) -> float:
        '''Return normalized score (0.0-1.0)'''
        pass

    def get_issues(self) -> List[Dict]:
        '''Return list of detected issues'''
        pass

Example usage:
from agents.analyzers.quality_analyzer import QualityAnalyzer

analyzer = QualityAnalyzer()
issues = analyzer.analyze('./rtl/design.v')
score = analyzer.get_score()"""

    doc.add_paragraph(analyzer_interface)

    doc.add_heading('7.4 Adapter Interface', level=2)

    adapter_interface = """All adapters implement common BaseStaticAdapter interface:

class BaseStaticAdapter:
    def execute(self, codebase_path: str) -> Dict[str, Any]:
        '''Execute tool and return structured results'''
        pass

    def parse_output(self, raw_output: str) -> Dict[str, Any]:
        '''Parse tool output into structured format'''
        pass

    def get_metrics(self) -> Dict[str, float]:
        '''Return normalized metrics'''
        pass

Example usage:
from agents.adapters.hierarchy_complexity_adapter import HierarchyComplexityAdapter

adapter = HierarchyComplexityAdapter()
results = adapter.execute('./rtl')
hierarchy = results['hierarchy']
complexity = results['complexity_metrics']"""

    doc.add_paragraph(adapter_interface)

    doc.add_page_break()

def add_output_formats(doc):
    """Add Output Formats section."""
    doc.add_heading('8. Output Formats', level=1)

    doc.add_heading('8.1 Design Health JSON', level=2)

    json_example = """The canonical designhealth.json output contains:

{
  "metadata": {
    "analysis_timestamp": "2026-02-27T12:00:00Z",
    "codebase_path": "./rtl",
    "version": "1.0"
  },
  "summary": {
    "total_files": 45,
    "total_lines": 12543,
    "total_modules": 23,
    "design_health_score": 78.5
  },
  "scores": {
    "synthesis_risk": 82,
    "lint_score": 75,
    "quality": 76,
    "maintainability": 79,
    "complexity": 71,
    "hierarchy": 85,
    "documentation": 68,
    "verification_coverage": 72
  },
  "issues": [
    {
      "type": "synthesis_safety",
      "severity": "HIGH",
      "file": "rtl/core.v",
      "line": 45,
      "message": "CDC violation detected",
      "suggestion": "Use gray code encoder/decoder for CDC"
    }
  ],
  "modules": [
    {
      "name": "core_module",
      "file": "rtl/core.v",
      "lines": 450,
      "complexity": 12,
      "quality_score": 82
    }
  ]
}"""

    doc.add_paragraph(json_example, style='List Paragraph')

    doc.add_heading('8.2 HTML Health Report', level=2)

    html_desc = """Interactive HTML report generated by healthreport_generator.py includes:

• Executive Summary: Overall design health with key metrics
• Scoring Breakdown: Visual charts of eight dimension scores
• Issue Summary: Categorized issues with severity levels
• Module Details: Per-module quality metrics and rankings
• Recommendations: LLM-generated improvement suggestions
• Trend Analysis: Historical quality trends (if applicable)
• Design Hierarchy: Interactive module dependency graph
• Synthesis Safety: CDC, combinational loop, and timing issues
• Verification Status: Testbench coverage and SVA usage"""

    doc.add_paragraph(html_desc)

    doc.add_heading('8.3 Excel Analysis Report', level=2)

    excel_desc = """Structured Excel workbook with sheets:

• Summary: Overall metrics and design health score
• Issues: Detailed issue list with location, severity, suggestions
• Modules: Per-module quality metrics and rankings
• Quality: Quality analyzer detailed results
• Complexity: Complexity metrics per module
• Synthesis Safety: Synthesis safety violations
• Hierarchy: Module instantiation hierarchy
• Analysis Metadata: Timestamps, configuration, tool versions"""

    doc.add_paragraph(excel_desc)

    doc.add_heading('8.4 NDJSON Export', level=2)

    ndjson_desc = """Newline-delimited JSON format for vector database ingestion:

Each line is a complete JSON object representing an analyzable chunk:

{"chunk_id": "core_v_0001", "module": "core", "type": "always_block", "content": "always @(...) begin...", "embedding": [...]}
{"chunk_id": "core_v_0002", "module": "core", "type": "module_instantiation", "content": "cpu_core inst (...)", "embedding": [...]}

Used for:
• Semantic search across codebase
• RAG-based context injection
• Similar code pattern detection
• Knowledge retention and reuse"""

    doc.add_paragraph(ndjson_desc)

    doc.add_heading('8.5 Natural Language Summary', level=2)

    summary_desc = """LLM-generated executive summaries include:

• Design Overview: High-level architecture assessment
• Key Findings: Top 3-5 most critical issues
• Quality Assessment: Strengths and weaknesses
• Recommendations: Prioritized improvement actions
• Verification Status: Testbench adequacy assessment
• Synthesis Readiness: Whether design is ready for synthesis
• Next Steps: Recommended follow-up actions

Example:
"This Verilog design implements a 32-bit processor core with good
modularity (score: 82) but has critical CDC violations in clock
domain crossing signals. The verification coverage is incomplete,
lacking assertions for pipeline hazard detection. Recommend:
1) Add gray code synchronizers for CDC signals
2) Implement SVA assertions for pipeline correctness
3) Increase testbench coverage to 95%+"""

    doc.add_paragraph(summary_desc)

    doc.add_page_break()

def add_extension_guide(doc):
    """Add Extension Guide section."""
    doc.add_heading('9. Extension Guide', level=1)

    doc.add_heading('9.1 Creating Custom Analyzers', level=2)

    custom_analyzer = """To create a custom analyzer, extend BaseAnalyzer:

from agents.analyzers.base_analyzer import BaseAnalyzer
from typing import Dict, Any, List

class CustomAnalyzer(BaseAnalyzer):
    def __init__(self):
        super().__init__()
        self.name = "custom_analyzer"

    def analyze(self, file_path: str) -> Dict[str, Any]:
        '''Analyze single file'''
        issues = []

        with open(file_path, 'r') as f:
            content = f.read()

        # Your analysis logic here
        if 'problematic_pattern' in content:
            issues.append({
                'severity': 'HIGH',
                'message': 'Problematic pattern found',
                'suggestion': 'Use recommended pattern instead'
            })

        return {
            'issues': issues,
            'score': 0.85  # 0.0-1.0
        }

    def get_score(self) -> float:
        '''Return normalized score'''
        return self._score if hasattr(self, '_score') else 1.0

# Register in MetricsCalculator
from agents.core.metrics_calculator import MetricsCalculator

class ExtendedMetricsCalculator(MetricsCalculator):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.custom_analyzer = CustomAnalyzer()"""

    doc.add_paragraph(custom_analyzer, style='List Paragraph')

    doc.add_heading('9.2 Creating Custom Adapters', level=2)

    custom_adapter = """To create a custom adapter, extend BaseStaticAdapter:

from agents.adapters.base_adapter import BaseStaticAdapter
import subprocess
from typing import Dict, Any

class CustomToolAdapter(BaseStaticAdapter):
    def __init__(self, debug: bool = False):
        super().__init__(debug=debug)
        self.tool_path = '/usr/bin/custom-tool'

    def execute(self, codebase_path: str) -> Dict[str, Any]:
        '''Execute external tool'''
        try:
            result = subprocess.run(
                [self.tool_path, codebase_path],
                capture_output=True,
                text=True,
                timeout=300
            )
            return self.parse_output(result.stdout)
        except Exception as e:
            self._logger.error(f"Tool execution failed: {e}")
            return {'error': str(e)}

    def parse_output(self, raw_output: str) -> Dict[str, Any]:
        '''Parse tool output'''
        # Parse and structure output
        metrics = {}
        for line in raw_output.split('\\n'):
            if 'key:value' in line:
                key, value = line.split(':')
                metrics[key] = float(value)
        return metrics

    def get_metrics(self) -> Dict[str, float]:
        '''Return normalized metrics'''
        return self._metrics if hasattr(self, '_metrics') else {}"""

    doc.add_paragraph(custom_adapter, style='List Paragraph')

    doc.add_heading('9.3 LLM Provider Integration', level=2)

    llm_provider = """To add a new LLM provider, extend LLMProvider interface:

from utils.common.llm_provider_base import LLMProviderBase

class CustomLLMProvider(LLMProviderBase):
    def __init__(self, api_key: str):
        self.api_key = api_key

    def generate(self, prompt: str, max_tokens: int = 16384) -> str:
        '''Generate response from prompt'''
        # Implement provider-specific API call
        response = self.client.generate(
            model='custom-model',
            prompt=prompt,
            max_tokens=max_tokens
        )
        return response.text

    def embed(self, text: str) -> List[float]:
        '''Generate embedding vector'''
        embedding = self.client.embed(
            model='custom-embed-model',
            text=text
        )
        return embedding.vector

# Register in utils/common/llm_tools.py
PROVIDERS = {
    'custom': CustomLLMProvider,
    'anthropic': AnthropicProvider,
    'qgenie': QGenieProvider,
}"""

    doc.add_paragraph(llm_provider, style='List Paragraph')

    doc.add_heading('9.4 Output Format Extensions', level=2)

    output_ext = """To add new output format, create custom generator:

from agents.parsers.base_report_generator import BaseReportGenerator

class CustomReportGenerator(BaseReportGenerator):
    def __init__(self, output_path: str):
        super().__init__(output_path)

    def generate(self, metrics: Dict[str, Any]) -> None:
        '''Generate report in custom format'''
        # Implement custom report generation
        report = self._build_report(metrics)

        with open(self.output_path, 'w') as f:
            f.write(report)

    def _build_report(self, metrics: Dict) -> str:
        '''Build custom report structure'''
        header = "=== Custom Report ===\\n"
        summary = f"Score: {metrics['design_health_score']}\\n"
        issues = self._format_issues(metrics['issues'])
        return header + summary + issues"""

    doc.add_paragraph(output_ext, style='List Paragraph')

    doc.add_page_break()

def add_dependencies(doc):
    """Add Dependencies & Technology Stack section."""
    doc.add_heading('10. Dependencies & Technology Stack', level=1)

    doc.add_heading('10.1 Python Packages', level=2)

    deps_text = """Core dependencies listed in requirements.txt:

langchain>=0.1.0          # LLM framework and RAG
langchain-anthropic>=0.1.0  # Anthropic Claude integration
anthropic>=0.7.0          # Anthropic SDK
sqlalchemy>=2.0           # Database ORM
psycopg2-binary>=2.9      # PostgreSQL driver
pgvector>=0.1.0           # pgvector Python client
python-docx>=0.8.11       # Word document generation
openpyxl>=3.0             # Excel file handling
pandas>=2.0               # Data manipulation
pyyaml>=6.0               # YAML configuration parsing
rich>=13.0                # Rich terminal output
requests>=2.28            # HTTP client
regex>=2023.0             # Advanced regex patterns
loguru>=0.6               # Enhanced logging
pydantic>=2.0             # Data validation
pydantic-settings>=2.0    # Configuration management"""

    doc.add_paragraph(deps_text, style='List Paragraph')

    doc.add_heading('10.2 External Tools', level=2)

    tools_table = doc.add_table(rows=6, cols=3)
    tools_table.style = 'Light Grid Accent 1'

    header_cells = tools_table.rows[0].cells
    header_cells[0].text = 'Tool'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'

    for cell in header_cells:
        set_cell_background(cell, '4F81BD')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    tools = [
        ('Verilator', '5.0+', 'Verilog static analysis and simulation'),
        ('Verible', '0.40+', 'SystemVerilog linting and formatting'),
        ('PostgreSQL', '14.0+', 'Primary database'),
        ('Pandoc', '2.19+', 'Document format conversion (optional)'),
        ('Mermaid CLI', '10.0+', 'Diagram rendering (optional)'),
    ]

    for i, (tool, version, purpose) in enumerate(tools, 1):
        row = tools_table.rows[i]
        row.cells[0].text = tool
        row.cells[1].text = version
        row.cells[2].text = purpose

    doc.add_heading('10.3 Database Stack', level=2)

    db_text = """PostgreSQL 14.0+
  - Primary relational database
  - Stores analysis results, module hierarchy, metrics

pgvector Extension
  - Vector database layer on top of PostgreSQL
  - Enables semantic similarity search
  - Supports L2, cosine, and inner product distances

LangChain Integrations
  - PostgreSQL vector store for embeddings
  - RAG pipeline support
  - Query augmentation

Connection Pool
  - SQLAlchemy connection pooling (5 persistent connections)
  - Automatic connection recycling (3600 seconds)
  - Health checks for stale connections"""

    doc.add_paragraph(db_text)

    doc.add_heading('10.4 LLM Providers', level=2)

    llm_table = doc.add_table(rows=5, cols=3)
    llm_table.style = 'Light Grid Accent 1'

    header_cells = llm_table.rows[0].cells
    header_cells[0].text = 'Provider'
    header_cells[1].text = 'Models'
    header_cells[2].text = 'Configuration'

    for cell in header_cells:
        set_cell_background(cell, '4F81BD')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    providers = [
        ('Anthropic', 'Claude Sonnet, Opus', 'Via anthropic SDK, API key required'),
        ('QGenie', 'Qwen, LLaMA', 'Local/cloud deployment'),
        ('Google Vertex AI', 'Gemini 2.5', 'Enterprise GCP integration'),
        ('Azure OpenAI', 'GPT-4, GPT-3.5', 'Azure subscription required'),
    ]

    for i, (provider, models, config) in enumerate(providers, 1):
        row = llm_table.rows[i]
        row.cells[0].text = provider
        row.cells[1].text = models
        row.cells[2].text = config

    doc.add_heading('10.5 Development Stack', level=2)

    dev_text = """Testing:
  - pytest: Unit and integration testing
  - pytest-cov: Code coverage measurement
  - pytest-mock: Mocking for isolated tests

Code Quality:
  - black: Python code formatting
  - pylint: Code analysis
  - mypy: Static type checking
  - isort: Import sorting

Documentation:
  - Sphinx: Documentation generation
  - doctest: Documentation testing

Version Control:
  - Git: Source control
  - GitHub/GitLab: Repository hosting"""

    doc.add_paragraph(dev_text)

    doc.add_heading('10.6 System Requirements', level=2)

    system_text = """Minimum:
  - Python 3.9+
  - 8GB RAM
  - 10GB disk space
  - PostgreSQL 14.0+

Recommended:
  - Python 3.11+
  - 16GB RAM
  - 50GB disk space
  - PostgreSQL 15.0+
  - Multi-core processor (4+ cores)

Supported Platforms:
  - Linux (Ubuntu 20.04+, RHEL 8+, Rocky 8+)
  - macOS (11+, both Intel and Apple Silicon)
  - Windows 11 with WSL2"""

    doc.add_paragraph(system_text)

def main():
    """Generate the comprehensive CARE Design Document."""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add all sections
    add_title_page(doc)
    add_table_of_contents(doc)
    add_executive_summary(doc)
    add_requirements_specification(doc)
    add_architecture_design(doc)
    add_configuration_reference(doc)
    add_project_structure(doc)
    add_usage_guide(doc)
    add_api_specification(doc)
    add_output_formats(doc)
    add_extension_guide(doc)
    add_dependencies(doc)

    # Save document
    output_path = '/sessions/nice-serene-darwin/mnt/CARE/CARE_Design_Document.docx'
    doc.save(output_path)

    print(f"Successfully generated CARE Design Document: {output_path}")
    print(f"Document includes:")
    print("  - Title page with version and date")
    print("  - Table of contents")
    print("  - 10 major sections with comprehensive content")
    print("  - Multiple tables with technical specifications")
    print("  - Configuration references and usage examples")
    print("  - API specifications and output formats")
    print(f"  - Total: ~20+ pages of professional documentation")

if __name__ == '__main__':
    main()
